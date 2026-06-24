# -*- coding: utf-8 -*-
import collections
import csv
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(r"D:\UE project\BlackDungeonRMK")
ASSET_ROOT = PROJECT_ROOT / "Content" / "Dungeon_Forge"
NEWWORLD = PROJECT_ROOT / "Content" / "Level" / "NewWorld.umap"
DOCS_DIR = PROJECT_ROOT / "Docs"
DATA_PATH = DOCS_DIR / "DungeonForge_asset_inventory_data.json"
DOCX_PATH = DOCS_DIR / "DungeonForge_NewWorld_Asset_Inventory.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 46)
MUTED = RGBColor(96, 108, 122)
BORDER = "D9E1EA"
HEADER_FILL = "E8EEF5"
GREEN_FILL = "EAF5EA"
AMBER_FILL = "FFF4D8"
GRAY_FILL = "F2F4F7"
FONT_EN = "Calibri"
FONT_CN = "Microsoft YaHei"


def extract_newworld_refs():
    data = NEWWORLD.read_bytes()
    refs = set()

    for match in re.finditer(rb"/Game/Dungeon_Forge[\x20-\x7e]{0,360}", data):
        text = match.group(0).decode("latin1", errors="ignore")
        text = re.split(r"[^A-Za-z0-9_/\.\-]+", text)[0]
        if text:
            refs.add(text)

    wide = data.decode("utf-16le", errors="ignore")
    for match in re.finditer(r"/Game/Dungeon_Forge[\x20-\x7e]{0,360}", wide):
        text = match.group(0)
        text = re.split(r"[^A-Za-z0-9_/\.\-]+", text)[0]
        if text:
            refs.add(text)

    return refs


def normalize_ref(ref):
    if "." in ref:
        ref = ref.split(".", 1)[0]
    return ref.rstrip("/")


def classify_asset(top, rel, stem, ext):
    if top == "Meshes" and stem.startswith("SM_"):
        return "Static Mesh"
    if top == "Meshes" and stem.startswith("SK_"):
        return "Skeletal Mesh"
    if top == "Blueprints" and stem.startswith(("BPP_", "BP_")):
        return "Placeable Blueprint"
    if top == "Levels" and ext == ".umap":
        return "Level/Sublevel"
    if top == "Decals":
        return "Decal/Material Effect"
    if top == "VFX" and stem.startswith("NS_"):
        return "Niagara System"
    if top == "VFX" and stem.startswith("BP_"):
        return "VFX Blueprint"
    if top == "VFX" and stem.startswith("SM_"):
        return "VFX Mesh"
    if top == "VFX":
        return "VFX Support"
    if top == "Demo" and "/Effects/" in rel and stem.startswith("P_"):
        return "Particle System"
    if top == "Demo" and "/Effects/Mesh/" in rel and stem.startswith("SM_"):
        return "VFX Mesh"
    if top == "Demo" and "/Effects/" in rel:
        return "Demo Effect Support"
    if top == "Demo" and "/StarterContent/Particles/" in rel and stem.startswith("P_"):
        return "Particle System"
    if stem.startswith(("MI_", "M_")):
        return "Material/Instance"
    if stem.startswith(("T_", "HDR_")):
        return "Texture"
    return "Other"


def applicability_for(item):
    name = item["name"].lower()
    folder = item["folder"].lower()
    typ = item["type"]

    if item["direct"]:
        return "已直接引用", "NewWorld 二进制路径扫描命中该资产，可视为当前场景已使用或已依赖。"

    if typ in {"Static Mesh", "Placeable Blueprint", "Level/Sublevel"}:
        if any(
            key in name or key in folder
            for key in [
                "floor",
                "wall",
                "column",
                "vault",
                "stair",
                "arch",
                "blind_arch",
                "stone_constructor",
                "groin",
            ]
        ):
            return "推荐适用", "适合扩展 NewWorld 的地牢墙体、拱券、地面、柱体、楼梯或通道结构。"
        if any(
            key in name or key in folder
            for key in [
                "door",
                "cage",
                "debris",
                "torch",
                "candle",
                "web",
                "skull",
                "statue",
                "mud",
                "wood",
                "board",
                "root",
                "hay",
            ]
        ):
            return "可选适用", "适合作为局部叙事、封闭空间细节、遮挡、道具或氛围点缀。"
        return "备选", "当前未发现直接引用，可按具体区域需求补充。"

    if typ in {"Niagara System", "VFX Blueprint", "Particle System", "Decal/Material Effect"}:
        if any(
            key in name or key in folder
            for key in ["dust", "fog", "candle", "godray", "blood", "decal", "leak", "moss", "grime"]
        ):
            return "推荐适用", "适合封闭地牢环境的尘埃、雾卡、烛火、光束、污渍或战斗痕迹。"
        if any(key in name or key in folder for key in ["rain", "splash", "water", "fluid", "smoke", "fire"]):
            return "可选适用", "适合湿区、火焰、雨水或水面交互段落；需按场景区域选择。"
        return "备选", "当前未发现直接引用，可作为后续氛围或交互特效候选。"

    if typ in {"VFX Mesh", "VFX Support", "Demo Effect Support"}:
        return "支撑资源", "通常作为特效系统的材质、贴图或网格依赖，随 Niagara/粒子系统间接使用。"

    return "支撑资源", "材质、贴图或辅助资源，通常由模型或特效间接引用。"


def collect_inventory():
    refs = extract_newworld_refs()
    ref_bases = {normalize_ref(ref) for ref in refs}

    items = []
    for path in ASSET_ROOT.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in {".uasset", ".umap", ".fbx"}:
            continue

        rel = path.relative_to(ASSET_ROOT).as_posix()
        stem = path.stem
        virtual = "/Game/Dungeon_Forge/" + rel.rsplit(".", 1)[0]
        top = rel.split("/")[0]
        folder = "/".join(rel.split("/")[:-1])
        typ = classify_asset(top, rel, stem, path.suffix.lower())

        items.append(
            {
                "rel": rel,
                "name": stem,
                "virtual": virtual,
                "top": top,
                "folder": folder,
                "type": typ,
                "size": path.stat().st_size,
                "direct": virtual in ref_bases,
            }
        )

    virtual_to_item = {item["virtual"]: item for item in items}
    for ref_base in sorted(ref_bases):
        if ref_base in virtual_to_item:
            continue
        candidates = [item for item in items if item["virtual"].startswith(ref_base + "_")]
        if len(candidates) == 1:
            candidates[0]["direct"] = True

    for item in items:
        item["applicability"], item["note"] = applicability_for(item)

    models = [
        item
        for item in items
        if item["type"] in {"Static Mesh", "Skeletal Mesh", "Placeable Blueprint", "Level/Sublevel"}
    ]
    vfx = [
        item
        for item in items
        if item["type"]
        in {
            "Niagara System",
            "VFX Blueprint",
            "Particle System",
            "Decal/Material Effect",
            "VFX Mesh",
            "VFX Support",
            "Demo Effect Support",
        }
    ]

    order = {"已直接引用": 0, "推荐适用": 1, "可选适用": 2, "备选": 3, "支撑资源": 4}
    models.sort(key=lambda item: (order.get(item["applicability"], 9), item["folder"], item["name"]))
    vfx.sort(key=lambda item: (order.get(item["applicability"], 9), item["folder"], item["name"]))

    summary = {
        "project_root": str(PROJECT_ROOT),
        "asset_root": str(ASSET_ROOT),
        "content_browser_root": "/All/Game/Dungeon_Forge",
        "newworld": str(NEWWORLD),
        "newworld_virtual": "/Game/Level/NewWorld",
        "total_files_uasset_umap_fbx": len(items),
        "model_rows": len(models),
        "vfx_rows": len(vfx),
        "raw_ref_count": len(refs),
        "direct_model_count": sum(1 for item in models if item["direct"]),
        "direct_vfx_count": sum(1 for item in vfx if item["direct"]),
        "model_type_counts": dict(collections.Counter(item["type"] for item in models)),
        "vfx_type_counts": dict(collections.Counter(item["type"] for item in vfx)),
        "model_app_counts": dict(collections.Counter(item["applicability"] for item in models)),
        "vfx_app_counts": dict(collections.Counter(item["applicability"] for item in vfx)),
        "direct_refs_raw": sorted(refs),
    }

    DATA_PATH.write_text(
        json.dumps({"summary": summary, "models": models, "vfx": vfx}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    for key, rows in [("models", models), ("vfx", vfx)]:
        csv_path = DOCS_DIR / f"DungeonForge_{key}_inventory.csv"
        with csv_path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.DictWriter(
                handle,
                fieldnames=["applicability", "type", "name", "virtual", "folder", "rel", "note", "direct", "size"],
            )
            writer.writeheader()
            for row in rows:
                writer.writerow({field: row.get(field, "") for field in writer.fieldnames})

    return summary, models, vfx


def set_run_font(run, size=None, bold=None, italic=None, color=None, name=FONT_EN, east_asia=FONT_CN):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_para_format(paragraph, before=0, after=6, line=1.25, keep_with_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next


def add_para(doc, text="", size=11, bold=False, color=INK, before=0, after=6, line=1.25):
    paragraph = doc.add_paragraph()
    set_para_format(paragraph, before, after, line)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    if level == 1:
        set_para_format(paragraph, before=18, after=10, line=1.0, keep_with_next=True)
        size, color = 16, BLUE
    elif level == 2:
        set_para_format(paragraph, before=14, after=7, line=1.0, keep_with_next=True)
        size, color = 13, BLUE
    else:
        set_para_format(paragraph, before=10, after=5, line=1.0, keep_with_next=True)
        size, color = 12, DARK_BLUE
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, size=8.0, bold=False, color=INK, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_para_format(paragraph, before=0, after=0, line=1.12)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_widths(table, widths_in):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(widths_in) * 1440)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), BORDER)

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths, font_size=8.0):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], HEADER_FILL)
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            size=8.5,
            bold=True,
            color=DARK_BLUE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    fill_by_label = {
        "已直接引用": GREEN_FILL,
        "推荐适用": HEADER_FILL,
        "可选适用": AMBER_FILL,
        "备选": GRAY_FILL,
        "支撑资源": "F7F7F7",
    }
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx >= 2 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], value, size=font_size, align=align)
            if idx == 0 and str(value) in fill_by_label:
                shade_cell(cells[idx], fill_by_label[str(value)])

    spacer = doc.add_paragraph()
    set_para_format(spacer, before=2, after=8, line=1.0)
    return table


def short_path(virtual_path):
    return virtual_path.replace("/Game/Dungeon_Forge/", "")


def asset_rows(assets):
    return [[item["applicability"], item["type"], item["name"], short_path(item["virtual"])] for item in assets]


def build_docx(summary, models, vfx):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_format(header, after=0, line=1.0)
    set_run_font(header.add_run("Dungeon_Forge Asset Inventory | NewWorld"), size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_format(footer, after=0, line=1.0)
    set_run_font(footer.add_run("BlackDungeonRMK - asset review"), size=9, color=MUTED)

    title = doc.add_paragraph()
    set_para_format(title, before=6, after=4, line=1.0)
    set_run_font(title.add_run("Dungeon_Forge 模型与特效资产清单"), size=22, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    set_para_format(subtitle, before=0, after=14, line=1.15)
    set_run_font(subtitle.add_run("面向 NewWorld 场景的适用性标注与引用证据"), size=12.5, color=MUTED)

    metadata_rows = [
        ("Content Browser 根路径", summary["content_browser_root"]),
        ("项目目录", summary["project_root"]),
        ("场景", "Content/Level/NewWorld.umap"),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("判定方式", "只读扫描 NewWorld.umap 中 /Game/Dungeon_Forge 路径引用；未直接引用的资产按地牢场景语境标注推荐适用性。"),
    ]
    metadata = doc.add_table(rows=0, cols=2)
    set_table_widths(metadata, [1.65, 4.85])
    for label, value in metadata_rows:
        cells = metadata.add_row().cells
        shade_cell(cells[0], HEADER_FILL)
        set_cell_text(cells[0], label, size=8.8, bold=True, color=DARK_BLUE)
        set_cell_text(cells[1], value, size=8.8, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "1. 总览", 1)
    model_app = collections.Counter(item["applicability"] for item in models)
    vfx_app = collections.Counter(item["applicability"] for item in vfx)
    add_table(
        doc,
        ["类别", "数量", "说明"],
        [
            ["总文件", summary["total_files_uasset_umap_fbx"], "Dungeon_Forge 下 .uasset / .umap / .fbx 文件总量"],
            [
                "模型/蓝图/关卡",
                summary["model_rows"],
                f"已直接引用 {summary['direct_model_count']}；推荐适用 {model_app.get('推荐适用', 0)}；可选适用 {model_app.get('可选适用', 0)}；备选 {model_app.get('备选', 0)}",
            ],
            [
                "特效/贴花/粒子",
                summary["vfx_rows"],
                f"直接引用 {summary['direct_vfx_count']}；推荐适用 {vfx_app.get('推荐适用', 0)}；可选适用 {vfx_app.get('可选适用', 0)}；支撑资源 {vfx_app.get('支撑资源', 0)}",
            ],
            [
                "NewWorld 路径引用",
                summary["raw_ref_count"],
                "扫描到的 /Game/Dungeon_Forge 原始引用条目数；类后缀与对象后缀会合并到对应资产。",
            ],
        ],
        [1.3, 0.75, 4.45],
        font_size=9.0,
    )

    add_heading(doc, "2. 标记规则", 1)
    add_table(
        doc,
        ["标记", "含义"],
        [
            ["已直接引用", "NewWorld.umap 中扫描到对应 /Game/Dungeon_Forge 路径，表示当前场景已使用或依赖。"],
            ["推荐适用", "未直接引用，但与 NewWorld 的封闭地牢结构、空气氛围、烛火/光束/污渍等强相关。"],
            ["可选适用", "适合局部叙事、道具点缀、湿区/火焰/水体等特定段落。"],
            ["备选", "可按后续关卡需求补充，当前未发现直接引用或高相关用途。"],
            ["支撑资源", "材质、贴图、VFX 网格等，通常作为模型或特效系统的间接依赖。"],
        ],
        [1.15, 5.35],
        font_size=9.0,
    )

    add_heading(doc, "3. NewWorld 已直接引用的 Dungeon_Forge 资产", 1)
    direct_models = [item for item in models if item["direct"]]
    add_table(
        doc,
        ["类型", "资产名", "Content Browser 路径"],
        [[item["type"], item["name"], short_path(item["virtual"])] for item in direct_models],
        [1.2, 2.0, 3.3],
        font_size=8.3,
    )
    add_para(
        doc,
        "未在 NewWorld.umap 中发现直接引用的 Dungeon_Forge 特效系统；特效清单中的“推荐适用”主要依据封闭地牢环境的氛围需求给出。",
        size=10,
        color=MUTED,
        after=8,
    )

    add_heading(doc, "4. NewWorld 使用建议", 1)
    add_para(
        doc,
        "结构扩展优先使用已直接引用和推荐适用的模块化蓝图/子关卡：ceiling_staircase、dungeon_column1、dungeon_floor_01/04、dungeon_random_mansory_01、dungeon_wall_big_1、groin_vault_family_des_1，以及 Blind_Arch 与 Mud 静态网格。",
        size=10.5,
    )
    add_para(
        doc,
        "氛围层建议优先加入 Dust_System / NS_Dust_System、NS_Fogcard、NS_Candle_01/02、BP_Godray、MI_Decal_Leaks、MI_DustPatch_Dark、血迹/污渍贴花等；这些资产能增强封闭地牢的空气感、潮湿感、战斗痕迹和光束层次。",
        size=10.5,
    )
    add_para(
        doc,
        "Demo/Effects 中的 Rain、Splash、Water、Fire 系列更适合特定空间段落，例如积水、漏水、火盆、湿区或触发式交互，不建议无差别铺满全场。",
        size=10.5,
    )

    add_heading(doc, "附录 A：模型 / 可放置蓝图 / 子关卡资产列表", 1)
    add_para(doc, "以下列表按适用性排序；路径均相对于 /All/Game/Dungeon_Forge。", size=9.5, color=MUTED)
    add_table(doc, ["适用标记", "类型", "资产名", "相对路径"], asset_rows(models), [0.8, 1.1, 1.85, 2.75], font_size=7.4)

    add_heading(doc, "附录 B：特效 / 贴花 / 粒子 / 特效支撑资源列表", 1)
    add_para(
        doc,
        "特效表包含 VFX、Decals、Demo/Effects 与粒子系统相关资源；材质/贴图类仅在其属于特效目录或特效演示依赖时列入。",
        size=9.5,
        color=MUTED,
    )
    add_table(doc, ["适用标记", "类型", "资产名", "相对路径"], asset_rows(vfx), [0.8, 1.15, 1.85, 2.7], font_size=7.4)

    add_heading(doc, "附录 C：NewWorld 扫描到的原始引用", 1)
    add_table(doc, ["原始引用字符串"], [[ref] for ref in summary["direct_refs_raw"]], [6.5], font_size=8.0)

    doc.save(DOCX_PATH)
    return DOCX_PATH


def main():
    DOCS_DIR.mkdir(exist_ok=True)
    summary, models, vfx = collect_inventory()
    docx_path = build_docx(summary, models, vfx)
    print(docx_path)


if __name__ == "__main__":
    main()
